import io
import base64
from pathlib import Path


def extract(file_bytes: bytes, filename: str) -> dict:
    """Extrae texto e imágenes del documento según su extensión."""
    ext = filename.lower().split(".")[-1]

    if ext == "pdf":
        return extract_pdf(file_bytes)
    elif ext in ["docx", "doc"]:
        return extract_docx(file_bytes)
    elif ext in ["png", "jpg", "jpeg"]:
        return extract_image(file_bytes)
    else:
        return {"text": "", "images": []}


def extract_pdf(file_bytes: bytes) -> dict:
    """Extrae texto e imágenes de un PDF usando PyMuPDF."""
    try:
        import fitz  # PyMuPDF

        doc = fitz.open(stream=file_bytes, filetype="pdf")
        text_parts = []
        images_b64 = []

        for page_num, page in enumerate(doc):
            # Extraer texto
            page_text = page.get_text("text")
            if page_text.strip():
                text_parts.append(f"[Página {page_num + 1}]\n{page_text}")

            # Extraer imágenes embebidas
            if len(images_b64) < 5:  # máximo 5 imágenes
                for img_index, img in enumerate(page.get_images(full=True)):
                    if len(images_b64) >= 5:
                        break
                    xref = img[0]
                    base_image = doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    # Convertir a JPEG base64
                    img_b64 = _to_jpeg_base64(image_bytes)
                    if img_b64:
                        images_b64.append(img_b64)

        doc.close()
        return {
            "text": "\n\n".join(text_parts),
            "images": images_b64,
        }
    except Exception as e:
        return {"text": f"[Error extrayendo PDF: {str(e)}]", "images": []}


def extract_docx(file_bytes: bytes) -> dict:
    """Extrae texto e imágenes de un archivo DOCX usando python-docx."""
    try:
        import docx
        import zipfile

        # Cargar documento desde bytes
        doc = docx.Document(io.BytesIO(file_bytes))
        text_parts = []

        # Extraer texto de párrafos
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Extraer texto de tablas
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    text_parts.append(row_text)

        # Extraer imágenes del ZIP interno del DOCX
        images_b64 = []
        with zipfile.ZipFile(io.BytesIO(file_bytes)) as zf:
            for name in zf.namelist():
                if len(images_b64) >= 5:
                    break
                if name.startswith("word/media/") and any(
                    name.lower().endswith(ext)
                    for ext in [".png", ".jpg", ".jpeg", ".gif", ".bmp"]
                ):
                    img_bytes = zf.read(name)
                    img_b64 = _to_jpeg_base64(img_bytes)
                    if img_b64:
                        images_b64.append(img_b64)

        return {
            "text": "\n".join(text_parts),
            "images": images_b64,
        }
    except Exception as e:
        return {"text": f"[Error extrayendo DOCX: {str(e)}]", "images": []}


def extract_image(file_bytes: bytes) -> dict:
    """La imagen ES el documento — la convierte a JPEG base64."""
    img_b64 = _to_jpeg_base64(file_bytes)
    return {
        "text": "[El documento es una imagen. Analiza su contenido visual completo.]",
        "images": [img_b64] if img_b64 else [],
    }


def _to_jpeg_base64(image_bytes: bytes) -> str | None:
    """Convierte bytes de imagen a JPEG base64."""
    try:
        from PIL import Image

        img = Image.open(io.BytesIO(image_bytes))

        # Convertir a RGB si tiene transparencia (PNG con alpha, etc.)
        if img.mode in ("RGBA", "LA", "P"):
            background = Image.new("RGB", img.size, (255, 255, 255))
            if img.mode == "P":
                img = img.convert("RGBA")
            background.paste(img, mask=img.split()[-1] if img.mode == "RGBA" else None)
            img = background
        elif img.mode != "RGB":
            img = img.convert("RGB")

        # Redimensionar si es muy grande (Claude tiene límites)
        max_dim = 1568
        if img.width > max_dim or img.height > max_dim:
            img.thumbnail((max_dim, max_dim), Image.LANCZOS)

        # Guardar como JPEG
        buffer = io.BytesIO()
        img.save(buffer, format="JPEG", quality=85, optimize=True)
        buffer.seek(0)

        return base64.b64encode(buffer.read()).decode("utf-8")
    except Exception:
        return None
